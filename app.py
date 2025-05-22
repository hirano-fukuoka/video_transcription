import streamlit as st
import whisper
from docx import Document
import tempfile
import os
import ffmpeg
import pysubs2

st.title("動画の音声＋字幕 自動抽出＆Word出力アプリ")

uploaded_file = st.file_uploader(
    "動画ファイルをアップロード（mp4, mkv, mov, etc.）", type=["mp4", "mkv", "mov", "avi"]
)

def extract_subtitles(video_path):
    # 字幕抽出 (ffmpegを使ってsrt抽出→pysubs2で読み込む)
    srt_path = video_path + ".srt"
    try:
        (
            ffmpeg
            .input(video_path)
            .output(srt_path, map='0:s:0')  # 最初の字幕ストリームを抽出
            .run(overwrite_output=True, quiet=True)
        )
        subs = pysubs2.load(srt_path)
        text = "\n".join([line.text for line in subs])
        os.remove(srt_path)
        return text
    except Exception as e:
        return "（字幕が検出できませんでした）"

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[-1]) as tmpfile:
        tmpfile.write(uploaded_file.getbuffer())
        tmpfile_path = tmpfile.name

    # 1. Whisperで音声文字起こし
    st.info("音声文字起こし中...")
    model = whisper.load_model("base")
    result = model.transcribe(tmpfile_path, language=None)
    transcription_text = result["text"]
    st.success("音声文字起こし完了！")

    # 2. ffmpeg＋pysubs2で字幕ストリーム抽出
    st.info("字幕ストリーム抽出中...")
    subtitle_text = extract_subtitles(tmpfile_path)
    if "字幕が検出できません" not in subtitle_text:
        st.success("字幕抽出成功！")
    else:
        st.warning("字幕ストリームが検出できませんでした（ハードサブの可能性あり）")

    # 3. 結果表示
    st.subheader("音声文字起こし結果")
    st.write(transcription_text)
    st.subheader("字幕抽出結果")
    st.write(subtitle_text)

    # 4. Wordファイル作成
    doc = Document()
    doc.add_heading('音声文字起こし', level=1)
    doc.add_paragraph(transcription_text)
    doc.add_heading('動画内字幕', level=1)
    doc.add_paragraph(subtitle_text)
    word_path = tmpfile_path + ".docx"
    doc.save(word_path)
    with open(word_path, "rb") as word_file:
        st.download_button(
            label="Wordファイルをダウンロード",
            data=word_file,
            file_name="transcription_with_subtitles.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    os.remove(tmpfile_path)
    os.remove(word_path)
